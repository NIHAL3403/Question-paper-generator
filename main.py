from huggingface_hub import notebook_login

notebook_login()

# Import necessary libraries
from docx import Document
import fitz  # PyMuPDF
from transformers import pipeline

# Function to read text from a Word document
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to read text from a PDF document
def read_pdf(file_path):
    pdf_document = fitz.open(file_path)
    full_text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        full_text += page.get_text()
    return full_text

# Function to generate questions using Bloom's taxonomy verbs
def generate_questions(text, num_questions=5):
    generator = pipeline('text-generation', model='distilgpt2')
    bloom_verbs = ["list", "explain", "describe", "compare", "contrast", "analyze", "evaluate", "create"]
    questions = []
    for i in range(num_questions):
        main_question_prompt = f"Generate a question using Bloom's taxonomy verb '{bloom_verbs[i % len(bloom_verbs)]}' based on the text: {text[:500]}"
        main_question = generator(main_question_prompt, max_new_tokens=150, num_return_sequences=1, truncation=True, pad_token_id=50256)[0]['generated_text'].strip()

        sub_questions = []
        for sub_label in ['a)', 'b)']:
            sub_question_prompt = f"Generate a sub-question using Bloom's taxonomy verb '{bloom_verbs[i % len(bloom_verbs)]}' based on the text: {text[:500]}"
            # Use max_new_tokens instead of max_length
            sub_question = generator(sub_question_prompt, max_new_tokens=75, num_return_sequences=1, truncation=True, pad_token_id=50256)[0]['generated_text'].strip()
            sub_questions.append(f"{sub_label} {sub_question}")

        questions.append((main_question, sub_questions))

    return questions

# Function to format the questions into a question paper
def format_question_paper(questions):
    question_paper = "IA Question Paper for 50 marks\n\n"
    for i, (main_question, sub_questions) in enumerate(questions, start=1):
        question_paper += f"Q{i}. {main_question} (10 marks)\n"
        for sub_question in sub_questions:
            question_paper += f"{sub_question}\n"
        question_paper += "\n"
    return question_paper

# Main function to integrate everything
def main(input_file):
    if input_file.endswith('.docx'):
        text = read_docx(input_file)
    elif input_file.endswith('.pdf'):
        text = read_pdf(input_file)
    else:
        raise ValueError("Unsupported file format. Please use a .docx or .pdf file.")

    questions = generate_questions(text)
    question_paper = format_question_paper(questions)

    with open('IA_Question_Paper.txt', 'w') as file:
        file.write(question_paper)
    print("Question paper generated successfully!")

# Upload your input file to Colab and set the path here
input_file = 'QA-1.docx'  # Replace with your input file path
main(input_file)

# Import necessary libraries
from docx import Document
import fitz  # PyMuPDF
from transformers import pipeline

# Function to read text from a Word document
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to read text from a PDF document
def read_pdf(file_path):
    pdf_document = fitz.open(file_path)
    full_text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        full_text += page.get_text()
    return full_text

# Function to generate questions using Bloom's taxonomy verbs
def generate_questions(text):
    generator = pipeline('text-generation', model='distilgpt2')
    bloom_verbs = ["list", "explain", "describe", "compare", "contrast", "analyze", "evaluate", "create"]
    questions = []
    for i in range(5):  # Generate 5 questions
        main_question_prompt = f"Generate a question using Bloom's taxonomy verb '{bloom_verbs[i % len(bloom_verbs)]}' based on the text: {text[:500]}"
        main_question = generator(main_question_prompt, max_new_tokens=150, num_return_sequences=1, truncation=True, pad_token_id=50256)[0]['generated_text'].strip()

        sub_questions = []
        for sub_label in ['a)', 'b)']:
            sub_question_prompt = f"Generate a sub-question using Bloom's taxonomy verb '{bloom_verbs[i % len(bloom_verbs)]}' based on the text: {text[:500]}"
            sub_question = generator(sub_question_prompt, max_new_tokens=75, num_return_sequences=1, truncation=True, pad_token_id=50256)[0]['generated_text'].strip()
            sub_questions.append(f"{sub_label} {sub_question}")

        questions.append((main_question, sub_questions))

    return questions

# Function to format the questions into a question paper
def format_question_paper(questions):
    question_paper = "IA Question Paper for 50 marks\n\n"
    for i, (main_question, sub_questions) in enumerate(questions, start=1):
        question_paper += f"Q{i}. {main_question} (10 marks)\n"
        for sub_question in sub_questions:
            question_paper += f"{sub_question}\n"
        question_paper += "\n"
    return question_paper

# Main function to integrate everything
def main(input_file):
    if input_file.endswith('.docx'):
        text = read_docx(input_file)
    elif input_file.endswith('.pdf'):
        text = read_pdf(input_file)
    else:
        raise ValueError("Unsupported file format. Please use a.docx or.pdf file.")

    questions = generate_questions(text)
    question_paper = format_question_paper(questions)

    with open('IA_Question_Paper.txt', 'w') as file:
        file.write(question_paper)
    print("Question paper generated successfully!")

# Upload your input file to Colab and set the path here
input_file = 'QA-1.docx'  # Replace with your input file path
main(input_file)